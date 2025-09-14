"""
DOCX Utilities Module for Medical Data Extraction Tool
Handles DOCX creation and formatting for medical reports with summary page
"""

import io
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import re


def create_a3_landscape_document():
    """Create a new A3 landscape document"""
    doc = Document()
    
    # Get the section (page setup)
    section = doc.sections[0]
    
    # Set A3 size in landscape (420mm x 297mm)
    section.page_width = Mm(420)   # Width in landscape
    section.page_height = Mm(297)  # Height in landscape
    section.orientation = WD_ORIENT.LANDSCAPE
    
    # Set margins
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    
    return doc


def add_summary_page(doc, summary_text):
    """Add a dedicated summary page to the document"""
    # Title
    title = doc.add_heading('MEDICAL DOCUMENT SUMMARY', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    # Add some space
    doc.add_paragraph()
    
    # Add summary content
    parse_markdown_to_docx(doc, summary_text, is_summary=True)
    
    # Add timestamp
    doc.add_paragraph()
    timestamp_para = doc.add_paragraph()
    timestamp_run = timestamp_para.add_run(
        f"Summary generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    timestamp_run.font.size = Pt(9)
    timestamp_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
    timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Page break before detailed content
    doc.add_page_break()


def parse_markdown_to_docx(doc, markdown_text, is_summary=False):
    """Parse markdown text and add to docx document"""
    lines = markdown_text.split('\n')
    
    # Set font sizes based on whether this is summary or detail
    if is_summary:
        header_sizes = [20, 18, 16, 14]  # Larger for summary
        normal_size = 12
    else:
        header_sizes = [18, 16, 14, 13]   # Smaller for detailed content
        normal_size = 12
    
    for line in lines:
        line = line.strip()
        
        if not line:  # Empty line
            doc.add_paragraph()
            continue
            
        # Headers
        if line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
            p.runs[0].font.size = Pt(header_sizes[3])
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            p.runs[0].font.size = Pt(header_sizes[2])
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            p.runs[0].font.size = Pt(header_sizes[1])
        elif line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.runs[0].font.size = Pt(header_sizes[0])
            
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(normal_size)
        elif re.match(r'^\d+\. ', line):
            content = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(content, style='List Number')
            for run in p.runs:
                run.font.size = Pt(normal_size)
            
        # Regular paragraph
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line, normal_size)


def add_formatted_text(paragraph, text, font_size=10):
    """Add text with basic markdown formatting (bold, italic)"""
    # Handle inline formatting
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(font_size)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic text
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(font_size)
        elif part.startswith('`') and part.endswith('`'):
            # Inline code
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(font_size - 1)
        else:
            # Regular text
            run = paragraph.add_run(part)
            run.font.size = Pt(font_size)


def create_docx_report(images_and_texts, filename, summary_text=None):
    """Create a DOCX report with optional summary page and side-by-side layout (image left, text right)"""
    
    # Create A3 landscape document
    doc = create_a3_landscape_document()
    
    # Add summary page if provided
    if summary_text and summary_text.strip():
        add_summary_page(doc, summary_text)
    
    # Add detailed content for each page with side-by-side layout
    for i, (image, extracted_text, page_label) in enumerate(images_and_texts):
        if i > 0 or summary_text:  # Add page break between documents
            doc.add_page_break()
            
        # Page header
        header = doc.add_heading(f'Medical Data - {page_label}', level=1)
        header_run = header.runs[0]
        header_run.font.size = Pt(16)
        header_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Create table for side-by-side layout (2 columns)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        
        # Set column widths (A3 landscape has plenty of space)
        table.columns[0].width = Mm(180)  # Left column for image
        table.columns[1].width = Mm(180)  # Right column for text
        
        # Remove table borders for cleaner look
        from docx.oxml import OxmlElement
        from docx.oxml.ns import nsdecls, qn
        
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Create border element
                tcBorders = OxmlElement('w:tcBorders')
                
                # Set all borders to nil (no border)
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)
                
                tcPr.append(tcBorders)
        
        # Get the cells
        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)
        
        # Left cell - Add image
        try:
            # Convert PIL image to bytes
            img_buffer = io.BytesIO()
            if image.mode in ("RGBA", "P"):
                image = image.convert("RGB")
            image.save(img_buffer, format='JPEG', quality=85)
            img_buffer.seek(0)
            
            # Add image to left cell
            left_para = left_cell.paragraphs[0]
            left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Calculate image size to fit in cell (maintain aspect ratio)
            img_width, img_height = image.size
            max_width = Mm(160)  # Leave some margin in cell
            max_height = Mm(180)  # A3 landscape has good height
            
            aspect_ratio = img_width / img_height
            if aspect_ratio > (max_width.inches / max_height.inches):
                # Image is wider - fit to width
                display_width = max_width
                display_height = Inches(max_width.inches / aspect_ratio)
            else:
                # Image is taller - fit to height
                display_height = max_height
                display_width = Inches(max_height.inches * aspect_ratio)
            
            run = left_para.add_run()
            run.add_picture(img_buffer, width=display_width, height=display_height)
            
        except Exception as e:
            # If image fails, add error message
            left_para = left_cell.paragraphs[0]
            error_run = left_para.add_run(f"Image loading error: {str(e)}")
            error_run.font.size = Pt(10)
            error_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        
        # Right cell - Add extracted text
        right_cell.paragraphs[0].clear()  # Clear default paragraph
        
        # Parse markdown content for right cell
        parse_markdown_to_docx_in_cell(right_cell, extracted_text)
        
        # Add some spacing after the table
        doc.add_paragraph()
        
        # Add footer info
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(
            f"Extracted on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | {page_label}"
        )
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def parse_markdown_to_docx_in_cell(cell, markdown_text):
    """Parse markdown text and add to a specific table cell"""
    lines = markdown_text.split('\n')
    
    # Font sizes for detailed content in cell
    header_sizes = [12, 11, 10, 9]
    normal_size = 8
    
    # Clear the default paragraph
    if cell.paragraphs:
        cell.paragraphs[0].clear()
        current_para = cell.paragraphs[0]
    else:
        current_para = cell.add_paragraph()
    
    para_used = False
    
    for line in lines:
        line = line.strip()
        
        if not line:  # Empty line
            if para_used:
                cell.add_paragraph()
                current_para = cell.add_paragraph()
                para_used = False
            continue
            
        # Headers
        if line.startswith('#### '):
            if para_used:
                current_para = cell.add_paragraph()
            add_formatted_text(current_para, f"**{line[5:]}**", header_sizes[3])
            para_used = True
        elif line.startswith('### '):
            if para_used:
                current_para = cell.add_paragraph()
            add_formatted_text(current_para, f"**{line[4:]}**", header_sizes[2])
            para_used = True
        elif line.startswith('## '):
            if para_used:
                current_para = cell.add_paragraph()
            add_formatted_text(current_para, f"**{line[3:]}**", header_sizes[1])
            para_used = True
        elif line.startswith('# '):
            if para_used:
                current_para = cell.add_paragraph()
            add_formatted_text(current_para, f"**{line[2:]}**", header_sizes[0])
            para_used = True
            
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            if para_used:
                current_para = cell.add_paragraph()
            add_formatted_text(current_para, f"• {line[2:]}", normal_size)
            para_used = True
        elif re.match(r'^\d+\. ', line):
            content = re.sub(r'^\d+\. ', '', line)
            if para_used:
                current_para = cell.add_paragraph()
            add_formatted_text(current_para, f"• {content}", normal_size)
            para_used = True
            
        # Regular paragraph
        else:
            if para_used:
                current_para = cell.add_paragraph()
            add_formatted_text(current_para, line, normal_size)
            para_used = True